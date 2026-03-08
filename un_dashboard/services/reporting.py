"""Comprehensive report builders for PDF, Word, and Excel exports."""

from __future__ import annotations

import base64
import io
import logging
import re
import zipfile
from datetime import datetime
from pathlib import Path
from typing import Any

import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
import streamlit as st
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Image, PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from plotly.subplots import make_subplots

from un_dashboard.core.constants import TARGET_INTERVIEWS_PER_ORG, YES_VALUES as CORE_YES_VALUES
from un_dashboard.design import ThemeMode, chart_palette_for, chart_scale_for, style_plotly_figure
from un_dashboard.services.transforms import (
    build_daily_interviews,
    build_org_daily_activity,
    build_partner_province_matrix,
    find_column,
    format_percent,
    score_positive_rate,
)


YES_VALUES = set(CORE_YES_VALUES)
STRONG_THRESHOLD = 0.80
MODERATE_THRESHOLD = 0.60
MAX_REPORT_TABLE_ROWS = 30
PLOT_IMAGE_WIDTH = 1200
PLOT_IMAGE_HEIGHT = 700
WEEKDAY_ORDER = ["Monday", "Tuesday", "Wednesday", "Thursday", "Friday", "Saturday", "Sunday"]
MAX_WORD_CHART_SECTIONS = 40
DOCX_IMAGE_WIDTH_INCH = 6.4
DOCX_CONTROL_CHAR_RE = re.compile(r"[\x00-\x08\x0B\x0C\x0E-\x1F]")
PLOT_IMAGE_FALLBACK_PNG_BYTES = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO+X8ksAAAAASUVORK5CYII="
)

_logger = logging.getLogger(__name__)
_kaleido_chrome_install_attempted = False
_plotly_image_export_warning_logged = False


def _safe_series(df: pd.DataFrame, column: str | None) -> pd.Series:
    if df.empty or not column or column not in df.columns:
        return pd.Series(dtype="object")
    selected = df[column]
    if isinstance(selected, pd.DataFrame):
        if selected.shape[1] == 0:
            return pd.Series(dtype="object")
        return selected.iloc[:, 0]
    return selected


def _safe_slug(value: str) -> str:
    clean = re.sub(r"[^A-Za-z0-9._-]+", "_", str(value or "").strip())
    return clean.strip("_") or "report"


def _safe_sheet_name(value: str) -> str:
    cleaned = re.sub(r"[\\/*?:\[\]]+", " ", str(value or "").strip())
    cleaned = " ".join(cleaned.split())
    return cleaned[:31] or "Sheet"


def _coerce_text(value: Any) -> str:
    if pd.isna(value):
        return ""
    if isinstance(value, pd.Timestamp):
        text = value.isoformat(sep=" ")
    else:
        text = str(value)
    return DOCX_CONTROL_CHAR_RE.sub(" ", text)


def _docx_safe_text(value: Any) -> str:
    return DOCX_CONTROL_CHAR_RE.sub(" ", _coerce_text(value))


def _display_df(df: pd.DataFrame, limit: int = MAX_REPORT_TABLE_ROWS) -> pd.DataFrame:
    if df is None or df.empty:
        return pd.DataFrame()
    out = df.head(limit).copy()
    for col in out.columns:
        out[col] = out[col].map(_coerce_text)
    return out


def _summary_export_df(items: list[dict[str, str]]) -> pd.DataFrame:
    return pd.DataFrame(items, columns=["Metric", "Value", "Interpretation"])


def _positive_interpretation(value: float) -> str:
    if pd.isna(value):
        return "Not available in current scope."
    if value >= STRONG_THRESHOLD:
        return "Strong result."
    if value >= MODERATE_THRESHOLD:
        return "Moderate result; monitor for consistency."
    return "Needs attention."


def _completion_interpretation(value: float) -> str:
    if pd.isna(value):
        return "Target baseline is not available."
    if value >= 1:
        return "Target achieved or exceeded."
    if value >= STRONG_THRESHOLD:
        return "Close to target."
    if value >= MODERATE_THRESHOLD:
        return "Moderate completion."
    return "Significant gap remains."


def _coverage_interpretation(value: float) -> str:
    if pd.isna(value):
        return "No valid responses captured."
    if value >= STRONG_THRESHOLD:
        return "High data completeness."
    if value >= MODERATE_THRESHOLD:
        return "Acceptable completeness."
    return "Low completeness; validate source data."


def _scope_target(scope_kind: str, data: pd.DataFrame, progress: pd.DataFrame) -> int:
    if scope_kind == "organization":
        return TARGET_INTERVIEWS_PER_ORG

    if not progress.empty and "target" in progress.columns:
        relevant = progress.copy()
        if "org_code" in data.columns and not data.empty:
            valid_orgs = data["org_code"].dropna().astype(str).unique().tolist()
            scoped = relevant[relevant["org_code"].astype(str).isin(valid_orgs)]
            if not scoped.empty:
                relevant = scoped
        total_target = int(pd.to_numeric(relevant["target"], errors="coerce").fillna(0).sum())
        if total_target > 0:
            return total_target

    org_count = int(data["org_code"].dropna().astype(str).nunique()) if "org_code" in data.columns else 0
    return max(org_count, 1) * TARGET_INTERVIEWS_PER_ORG


def _date_window(data: pd.DataFrame, date_column: str | None) -> tuple[str, str, str]:
    if data.empty or not date_column or date_column not in data.columns:
        return "Not available", "Not available", "No valid interview dates found."

    parsed = pd.to_datetime(_safe_series(data, date_column), errors="coerce", utc=True)
    if isinstance(parsed, pd.DatetimeIndex):
        parsed = pd.Series(parsed, index=data.index)
    valid = parsed.dropna()
    if valid.empty:
        return "Not available", "Not available", "No valid interview dates found."

    start_date = valid.min().tz_convert(None).date()
    end_date = valid.max().tz_convert(None).date()
    label = f"{start_date.isoformat()} to {end_date.isoformat()}"
    return start_date.isoformat(), end_date.isoformat(), label


def _weekday_activity_table(data: pd.DataFrame, date_column: str | None) -> pd.DataFrame:
    if data.empty or not date_column or date_column not in data.columns:
        return pd.DataFrame(columns=["weekday_name", "interviews"])

    parsed = pd.to_datetime(_safe_series(data, date_column), errors="coerce", utc=True)
    if isinstance(parsed, pd.DatetimeIndex):
        parsed = pd.Series(parsed, index=data.index)
    valid = parsed.dropna()
    if valid.empty:
        return pd.DataFrame(columns=["weekday_name", "interviews"])

    out = (
        valid.dt.tz_convert(None)
        .dt.day_name()
        .value_counts()
        .rename_axis("weekday_name")
        .reset_index(name="interviews")
    )
    out["weekday_name"] = pd.Categorical(out["weekday_name"], categories=WEEKDAY_ORDER, ordered=True)
    return out.sort_values("weekday_name").reset_index(drop=True)


def _preferred_date_column(data: pd.DataFrame, indicators: dict[str, str | None]) -> str | None:
    start_column = find_column(data, ["start"])
    if start_column and start_column in data.columns:
        return start_column

    date_column = indicators.get("date")
    if date_column and date_column in data.columns:
        return date_column

    fallback = find_column(data, ["date_time", "submissiondate", "interview_date", "today"])
    if fallback and fallback in data.columns:
        return fallback
    return None


def _count_table(series: pd.Series, *, top_n: int = 12, title_col: str = "value") -> pd.DataFrame:
    if series.empty:
        return pd.DataFrame(columns=[title_col, "count"])
    cleaned = series.fillna("Unknown").astype(str).str.strip().replace("", "Unknown")
    return cleaned.value_counts(dropna=False).head(top_n).rename_axis(title_col).reset_index(name="count")


def _pick_distribution_column(data: pd.DataFrame, indicators: dict[str, str | None]) -> str | None:
    candidates: list[str | None] = [
        indicators.get("gender"),
        find_column(data, ["district", "District"]),
        "province" if "province" in data.columns else None,
        "ingo_partner" if "ingo_partner" in data.columns else None,
        "sheet_name" if "sheet_name" in data.columns else None,
    ]
    for column in candidates:
        if not column or column not in data.columns:
            continue
        distinct = int(_safe_series(data, column).dropna().astype(str).nunique())
        if 1 < distinct <= 20:
            return column
    return None


def _relevant_progress(scope_kind: str, data: pd.DataFrame, progress: pd.DataFrame) -> pd.DataFrame:
    if progress.empty:
        return progress.copy()
    if scope_kind == "organization":
        return progress.head(1).copy()

    out = progress.copy()
    if "org_code" in data.columns and not data.empty:
        valid_orgs = data["org_code"].dropna().astype(str).unique().tolist()
        scoped = out[out["org_code"].astype(str).isin(valid_orgs)]
        if not scoped.empty:
            out = scoped
    return out.sort_values(["interviews", "progress_pct"], ascending=[False, False]).reset_index(drop=True)


def _figure_to_png_bytes(fig) -> bytes:
    global _kaleido_chrome_install_attempted, _plotly_image_export_warning_logged

    fig.update_layout(width=PLOT_IMAGE_WIDTH, height=PLOT_IMAGE_HEIGHT)
    try:
        return fig.to_image(format="png", width=PLOT_IMAGE_WIDTH, height=PLOT_IMAGE_HEIGHT, scale=1)
    except Exception as exc:
        error_text = str(exc).lower()
        # Kaleido v1+ needs a browser; try one auto-install attempt on cloud runners.
        if "kaleido" in error_text and "chrome" in error_text and not _kaleido_chrome_install_attempted:
            _kaleido_chrome_install_attempted = True
            try:
                import kaleido

                install_chrome = getattr(kaleido, "get_chrome_sync", None)
                if callable(install_chrome):
                    install_chrome()
                    return fig.to_image(format="png", width=PLOT_IMAGE_WIDTH, height=PLOT_IMAGE_HEIGHT, scale=1)
            except Exception:
                pass

        if not _plotly_image_export_warning_logged:
            _plotly_image_export_warning_logged = True
            _logger.warning(
                "Plotly image export failed. Falling back to placeholder PNG for report charts. Error: %s",
                exc,
            )
        return PLOT_IMAGE_FALLBACK_PNG_BYTES


def _org_activity_chart(activity_df: pd.DataFrame, org_label: str, theme_mode: ThemeMode):
    palette = chart_palette_for(theme_mode)
    scale = chart_scale_for(theme_mode)
    template = "plotly_dark" if theme_mode == "dark" else "plotly_white"

    ordered = activity_df.sort_values("date").copy()
    ordered["date_label"] = ordered["date"].astype(str)
    ordered["weekday_name"] = pd.Categorical(ordered["weekday_name"], categories=WEEKDAY_ORDER, ordered=True)

    heat_grid = (
        ordered.pivot_table(
            index="weekday_name",
            columns="date_label",
            values="interviews",
            aggfunc="sum",
            fill_value=0,
            observed=False,
        )
        .reindex(WEEKDAY_ORDER)
        .fillna(0)
    )

    fig = make_subplots(
        rows=2,
        cols=1,
        shared_xaxes=True,
        vertical_spacing=0.09,
        row_heights=[0.72, 0.28],
        specs=[[{"secondary_y": True}], [{"secondary_y": False}]],
    )
    fig.add_trace(
        go.Bar(
            x=ordered["date_label"],
            y=ordered["interviews"],
            name="Daily interviews",
            marker=dict(
                color=ordered["interviews"],
                colorscale=scale,
                line=dict(color=palette[0], width=1),
            ),
            customdata=ordered[["weekday_name", "cumulative"]],
            hovertemplate="Date: %{x}<br>Day: %{customdata[0]}<br>Interviews: %{y}<br>Cumulative: %{customdata[1]}<extra></extra>",
        ),
        row=1,
        col=1,
        secondary_y=False,
    )
    fig.add_trace(
        go.Scatter(
            x=ordered["date_label"],
            y=ordered["cumulative"],
            name="Cumulative",
            mode="lines+markers",
            line=dict(color=palette[1], width=3),
            marker=dict(size=7, color=palette[1]),
            hovertemplate="Date: %{x}<br>Cumulative: %{y}<extra></extra>",
        ),
        row=1,
        col=1,
        secondary_y=True,
    )
    fig.add_trace(
        go.Heatmap(
            x=heat_grid.columns.tolist(),
            y=heat_grid.index.tolist(),
            z=heat_grid.values,
            name="Weekday pattern",
            colorscale=scale,
            text=heat_grid.values,
            texttemplate="%{text}",
            hovertemplate="Date: %{x}<br>Day: %{y}<br>Interviews: %{z}<extra></extra>",
            showscale=False,
        ),
        row=2,
        col=1,
    )
    fig.update_layout(
        title=f"{org_label} Interview Activity by Start Date",
        template=template,
        barmode="overlay",
        legend=dict(orientation="h"),
    )
    fig.update_yaxes(title_text="Interviews", row=1, col=1, secondary_y=False)
    fig.update_yaxes(title_text="Cumulative", row=1, col=1, secondary_y=True)
    fig.update_yaxes(title_text="Day of Week", row=2, col=1)
    fig.update_xaxes(title_text="Interview date", row=2, col=1, tickangle=-35)
    style_plotly_figure(fig, theme_mode)
    return fig


def _org_activity_sections(
    data: pd.DataFrame,
    date_column: str | None,
    theme_mode: ThemeMode,
) -> list[dict[str, Any]]:
    activity = build_org_daily_activity(data, date_column)
    if activity.empty:
        return []

    sections: list[dict[str, Any]] = []
    for org_code, org_df in activity.groupby("org_code", sort=True):
        ordered = org_df.sort_values("date").copy()
        preview = ordered[["date", "weekday_name", "interviews", "cumulative"]].copy()
        preview["date"] = preview["date"].astype(str)
        fig = _org_activity_chart(ordered, str(org_code), theme_mode)
        sections.append(
            {
                "org_code": str(org_code),
                "title": f"{org_code} start-date activity",
                "caption": "Daily interviews by calendar date using the `start` field, with weekday pattern and cumulative progress.",
                "image": _figure_to_png_bytes(fig),
                "table": preview,
            }
        )
    return sections


def _build_kpi_summary(
    scope_kind: str,
    scope_label: str,
    data: pd.DataFrame,
    progress: pd.DataFrame,
    indicators: dict[str, str | None],
) -> tuple[list[dict[str, str]], pd.DataFrame, dict[str, float]]:
    interviews = int(len(data))
    organizations = int(data["org_code"].dropna().astype(str).nunique()) if "org_code" in data.columns else 0
    projects = int(data["sheet_name"].dropna().astype(str).nunique()) if "sheet_name" in data.columns else organizations
    provinces = int(data["province"].dropna().astype(str).nunique()) if "province" in data.columns else 0
    partners = int(data["ingo_partner"].dropna().astype(str).nunique()) if "ingo_partner" in data.columns else 0

    target_total = _scope_target(scope_kind, data, progress)
    completion = interviews / target_total if target_total else float("nan")
    date_column = _preferred_date_column(data, indicators)
    trend = build_daily_interviews(data, date_column)
    active_days = int(len(trend))
    peak_daily = int(trend["interviews"].max()) if not trend.empty else 0
    avg_daily = float(interviews / active_days) if active_days else float("nan")

    training_col = indicators.get("training")
    helpful_col = indicators.get("helpful")
    complaint_col = indicators.get("complaint_awareness")

    training_rate = score_positive_rate(_safe_series(data, training_col), YES_VALUES)
    helpful_rate = score_positive_rate(_safe_series(data, helpful_col))
    complaint_rate = score_positive_rate(_safe_series(data, complaint_col), YES_VALUES)

    indicator_cols = [col for col in [training_col, helpful_col, complaint_col] if col and col in data.columns]
    if indicator_cols:
        completeness = float(
            pd.concat(
                [_safe_series(data, col).fillna("").astype(str).str.strip().ne("").rename(col) for col in indicator_cols],
                axis=1,
            ).mean(axis=1).mean()
        )
    else:
        completeness = float("nan")

    summary_items = [
        {"Metric": "Scope", "Value": scope_label, "Interpretation": "Current report scope."},
        {"Metric": "Interviews analyzed", "Value": f"{interviews:,}", "Interpretation": "Records included after active filters."},
        {"Metric": "Completion", "Value": format_percent(completion), "Interpretation": _completion_interpretation(completion)},
        {
            "Metric": "Organizations covered" if scope_kind == "public" else "Projects covered",
            "Value": f"{organizations:,}" if scope_kind == "public" else f"{projects:,}",
            "Interpretation": "Distinct entities represented in this report.",
        },
        {"Metric": "Provinces covered", "Value": f"{provinces:,}", "Interpretation": "Geographic spread in current scope."},
        {"Metric": "Partners covered", "Value": f"{partners:,}", "Interpretation": "Distinct INGO partners in current scope."},
        {"Metric": "Active interview days", "Value": f"{active_days:,}", "Interpretation": "Distinct dates with at least one interview."},
        {"Metric": "Peak daily interviews", "Value": f"{peak_daily:,}", "Interpretation": "Maximum interviews collected on a single day."},
        {"Metric": "Avg interviews per active day", "Value": "N/A" if pd.isna(avg_daily) else f"{avg_daily:.1f}", "Interpretation": "Average field throughput per active collection day."},
        {"Metric": "Training attendance", "Value": format_percent(training_rate), "Interpretation": _positive_interpretation(training_rate)},
        {"Metric": "Helpfulness", "Value": format_percent(helpful_rate), "Interpretation": _positive_interpretation(helpful_rate)},
        {"Metric": "Complaint awareness", "Value": format_percent(complaint_rate), "Interpretation": _positive_interpretation(complaint_rate)},
        {"Metric": "Data completeness", "Value": format_percent(completeness), "Interpretation": _coverage_interpretation(completeness)},
    ]
    indicator_df = pd.DataFrame(
        [
            {"Indicator": "Training attendance", "Positive rate": format_percent(training_rate), "Interpretation": _positive_interpretation(training_rate)},
            {"Indicator": "Helpfulness", "Positive rate": format_percent(helpful_rate), "Interpretation": _positive_interpretation(helpful_rate)},
            {"Indicator": "Complaint awareness", "Positive rate": format_percent(complaint_rate), "Interpretation": _positive_interpretation(complaint_rate)},
            {"Indicator": "Data completeness", "Positive rate": format_percent(completeness), "Interpretation": _coverage_interpretation(completeness)},
        ]
    )
    return summary_items, indicator_df, {
        "completion": completion,
        "training": training_rate,
        "helpfulness": helpful_rate,
        "complaint_awareness": complaint_rate,
        "completeness": completeness,
    }


def _build_insights(
    scope_kind: str,
    scope_label: str,
    data: pd.DataFrame,
    progress: pd.DataFrame,
    indicators: dict[str, str | None],
    scores: dict[str, float],
) -> tuple[list[str], list[str], list[str]]:
    date_column = _preferred_date_column(data, indicators)
    _, _, date_label = _date_window(data, date_column)
    insights = [
        f"The report covers {len(data):,} interviews for {scope_label}.",
        f"Interview date coverage: {date_label}.",
        f"Overall completion stands at {format_percent(scores['completion'])}.",
    ]

    if not data.empty and "province" in data.columns:
        top_province = _count_table(_safe_series(data, "province"), top_n=1, title_col="province")
        if not top_province.empty:
            row = top_province.iloc[0]
            insights.append(f"The highest interview concentration is in {row['province']} with {int(row['count']):,} interviews.")

    if scope_kind == "public" and not progress.empty:
        top_progress = progress.sort_values(["interviews", "progress_pct"], ascending=[False, False]).head(1)
        if not top_progress.empty:
            row = top_progress.iloc[0]
            insights.append(f"Top organization by volume is {row['org_code']} with {int(row['interviews']):,} interviews and {format_percent(float(row['progress_pct']))} completion.")

    trend = build_daily_interviews(data, date_column)
    if not trend.empty:
        peak_row = trend.sort_values("interviews", ascending=False).head(1).iloc[0]
        insights.append(f"Peak daily collection was {int(peak_row['interviews']):,} interviews on {peak_row['date']}.")

    risks: list[str] = []
    if not pd.isna(scores["completion"]) and scores["completion"] < MODERATE_THRESHOLD:
        risks.append("Completion remains below 60%; field collection should be accelerated.")
    if not pd.isna(scores["complaint_awareness"]) and scores["complaint_awareness"] < MODERATE_THRESHOLD:
        risks.append("Complaint awareness is low; safeguarding and accountability messaging needs reinforcement.")
    if not pd.isna(scores["helpfulness"]) and scores["helpfulness"] < MODERATE_THRESHOLD:
        risks.append("Perceived helpfulness is below the desired threshold; review service relevance and follow-up support.")
    if not pd.isna(scores["completeness"]) and scores["completeness"] < MODERATE_THRESHOLD:
        risks.append("Data completeness is low; validate collection and correction workflows.")
    if not risks:
        risks.append("No critical performance risk was detected in the current filtered scope.")

    recommendations: list[str] = []
    if not pd.isna(scores["completion"]) and scores["completion"] < 1:
        recommendations.append("Prioritize remaining organizations or respondent groups with the largest completion gaps.")
    if not pd.isna(scores["training"]) and scores["training"] < STRONG_THRESHOLD:
        recommendations.append("Strengthen training follow-up and attendance tracking for partially reached beneficiaries.")
    if not pd.isna(scores["complaint_awareness"]) and scores["complaint_awareness"] < STRONG_THRESHOLD:
        recommendations.append("Embed complaint-contact reminders in every interaction and monitoring visit.")
    if not pd.isna(scores["helpfulness"]) and scores["helpfulness"] < STRONG_THRESHOLD:
        recommendations.append("Use qualitative follow-up to identify why support is not translating into perceived usefulness.")
    if not recommendations:
        recommendations.append("Maintain current delivery quality and continue monitoring for consistency across locations.")

    return insights, risks, recommendations


def _public_chart_sections(
    data: pd.DataFrame,
    progress: pd.DataFrame,
    indicators: dict[str, str | None],
    theme_mode: ThemeMode,
    scores: dict[str, float],
) -> list[dict[str, Any]]:
    sections: list[dict[str, Any]] = []
    palette = chart_palette_for(theme_mode)
    scale = chart_scale_for(theme_mode)
    template = "plotly_dark" if theme_mode == "dark" else "plotly_white"
    date_column = _preferred_date_column(data, indicators)

    if not progress.empty:
        fig = px.bar(
            progress.sort_values("interviews", ascending=False),
            x="org_code",
            y="interviews",
            color="status" if "status" in progress.columns else "ingo_partner",
            template=template,
            title="Organization Progress",
            color_discrete_sequence=palette,
        )
        fig.add_hline(y=TARGET_INTERVIEWS_PER_ORG, line_dash="dash", line_color=palette[1])
        fig.update_layout(xaxis_title="Organization", yaxis_title="Interviews")
        style_plotly_figure(fig, theme_mode)
        sections.append(
            {
                "title": "Organization progress",
                "caption": "Interview volume by organization against the per-organization target.",
                "image": _figure_to_png_bytes(fig),
                "table": progress[["org_code", "province", "ingo_partner", "interviews", "target", "progress_pct", "status"]].copy(),
            }
        )

    trend = build_daily_interviews(data, date_column)
    if not trend.empty:
        fig = go.Figure()
        fig.add_trace(go.Bar(x=trend["date"], y=trend["interviews"], name="Daily interviews", marker_color=palette[0]))
        fig.add_trace(go.Scatter(x=trend["date"], y=trend["cumulative"], name="Cumulative", mode="lines+markers", line=dict(color=palette[1], width=3)))
        fig.update_layout(title="Interview Collection Trend", xaxis_title="Date", yaxis_title="Interviews", template=template)
        style_plotly_figure(fig, theme_mode)
        sections.append({"title": "Collection trend", "caption": "Daily and cumulative interview collection over time.", "image": _figure_to_png_bytes(fig), "table": trend.copy()})

    weekday = _weekday_activity_table(data, date_column)
    if not weekday.empty:
        fig = px.bar(
            weekday,
            x="weekday_name",
            y="interviews",
            template=template,
            title="Weekly Collection Pattern",
            color="interviews",
            color_continuous_scale=scale,
        )
        fig.update_layout(showlegend=False, xaxis_title="Weekday", yaxis_title="Interviews")
        style_plotly_figure(fig, theme_mode)
        sections.append(
            {
                "title": "Weekly activity pattern",
                "caption": "Interview volume aggregated by weekday using the start/date field.",
                "image": _figure_to_png_bytes(fig),
                "table": weekday.copy(),
            }
        )

    org_activity = build_org_daily_activity(data, date_column)
    if not org_activity.empty:
        timeline = org_activity.copy()
        timeline["date"] = timeline["date"].astype(str)
        fig = px.density_heatmap(
            timeline,
            x="date",
            y="org_code",
            z="interviews",
            histfunc="sum",
            template=template,
            color_continuous_scale=scale,
            title="Organization Interview Calendar",
            text_auto=True,
        )
        fig.update_layout(xaxis_title="Date", yaxis_title="Organization")
        style_plotly_figure(fig, theme_mode)
        sections.append(
            {
                "title": "Organization interview calendar",
                "caption": "Heatmap of daily interview counts per organization based on the `start` field.",
                "image": _figure_to_png_bytes(fig),
                "table": timeline[["org_code", "date", "weekday_name", "interviews", "cumulative"]].copy(),
            }
        )

    matrix = build_partner_province_matrix(data)
    if not matrix.empty:
        melted = matrix.melt(id_vars=["ingo_partner"], var_name="province", value_name="interviews")
        fig = px.density_heatmap(
            melted,
            x="province",
            y="ingo_partner",
            z="interviews",
            histfunc="sum",
            template=template,
            color_continuous_scale=scale,
            title="Partner x Province Coverage",
        )
        fig.update_layout(xaxis_title="Province", yaxis_title="Partner")
        style_plotly_figure(fig, theme_mode)
        sections.append({"title": "Partner and province coverage", "caption": "Heatmap of interview concentration by partner and province.", "image": _figure_to_png_bytes(fig), "table": melted.copy()})

    score_df = pd.DataFrame(
        [
            {"indicator": "Completion", "score": scores["completion"]},
            {"indicator": "Training", "score": scores["training"]},
            {"indicator": "Helpfulness", "score": scores["helpfulness"]},
            {"indicator": "Complaint awareness", "score": scores["complaint_awareness"]},
            {"indicator": "Data completeness", "score": scores["completeness"]},
        ]
    ).dropna(subset=["score"])
    if not score_df.empty:
        fig = px.bar(
            score_df,
            x="indicator",
            y="score",
            template=template,
            title="Analytical Scorecard",
            color="score",
            color_continuous_scale=scale,
            text=score_df["score"].map(lambda x: f"{x * 100:.1f}%"),
        )
        fig.update_yaxes(tickformat=".0%", range=[0, 1])
        fig.update_layout(showlegend=False, xaxis_title="Indicator", yaxis_title="Score")
        style_plotly_figure(fig, theme_mode)
        sections.append({"title": "Analytical scorecard", "caption": "High-level indicator scores for the current public report scope.", "image": _figure_to_png_bytes(fig), "table": score_df.assign(score=score_df["score"].map(lambda x: f"{x * 100:.1f}%"))})

    return sections


def _organization_chart_sections(
    scope_label: str,
    data: pd.DataFrame,
    progress: pd.DataFrame,
    indicators: dict[str, str | None],
    theme_mode: ThemeMode,
    scores: dict[str, float],
) -> list[dict[str, Any]]:
    sections: list[dict[str, Any]] = []
    palette = chart_palette_for(theme_mode)
    scale = chart_scale_for(theme_mode)
    template = "plotly_dark" if theme_mode == "dark" else "plotly_white"
    date_column = _preferred_date_column(data, indicators)

    progress_value = float(progress["interviews"].iloc[0]) if not progress.empty and "interviews" in progress.columns else float(len(data))
    progress_pct = progress_value / TARGET_INTERVIEWS_PER_ORG if TARGET_INTERVIEWS_PER_ORG else float("nan")
    target_df = pd.DataFrame(
        [
            {"metric": "Completed interviews", "value": progress_value},
            {"metric": "Remaining to target", "value": max(TARGET_INTERVIEWS_PER_ORG - progress_value, 0)},
        ]
    )
    fig = px.bar(target_df, x="metric", y="value", color="metric", template=template, title=f"{scope_label} Progress vs Target", color_discrete_sequence=palette)
    fig.update_layout(showlegend=False, xaxis_title="", yaxis_title="Interviews")
    style_plotly_figure(fig, theme_mode)
    sections.append(
        {
            "title": "Progress vs target",
            "caption": f"Current organization progress is {format_percent(progress_pct)} of the {TARGET_INTERVIEWS_PER_ORG}-interview target.",
            "image": _figure_to_png_bytes(fig),
            "table": pd.DataFrame([{"Metric": "Completed interviews", "Value": int(progress_value)}, {"Metric": "Remaining interviews", "Value": int(max(TARGET_INTERVIEWS_PER_ORG - progress_value, 0))}, {"Metric": "Completion", "Value": format_percent(progress_pct)}]),
        }
    )

    trend = build_daily_interviews(data, date_column)
    if not trend.empty:
        fig = go.Figure()
        fig.add_trace(go.Bar(x=trend["date"], y=trend["interviews"], name="Daily interviews", marker_color=palette[0]))
        fig.add_trace(go.Scatter(x=trend["date"], y=trend["cumulative"], name="Cumulative", mode="lines+markers", line=dict(color=palette[1], width=3)))
        fig.update_layout(title=f"{scope_label} Collection Trend", xaxis_title="Date", yaxis_title="Interviews", template=template)
        style_plotly_figure(fig, theme_mode)
        sections.append({"title": "Collection trend", "caption": "Daily and cumulative interview collection for the selected organization.", "image": _figure_to_png_bytes(fig), "table": trend.copy()})

    weekday = _weekday_activity_table(data, date_column)
    if not weekday.empty:
        fig = px.bar(
            weekday,
            x="weekday_name",
            y="interviews",
            template=template,
            title=f"{scope_label} Weekly Collection Pattern",
            color="interviews",
            color_continuous_scale=scale,
        )
        fig.update_layout(showlegend=False, xaxis_title="Weekday", yaxis_title="Interviews")
        style_plotly_figure(fig, theme_mode)
        sections.append(
            {
                "title": "Weekly activity pattern",
                "caption": "Interview volume aggregated by weekday from the organization start/date column.",
                "image": _figure_to_png_bytes(fig),
                "table": weekday.copy(),
            }
        )

    score_df = pd.DataFrame(
        [
            {"indicator": "Training", "score": scores["training"]},
            {"indicator": "Helpfulness", "score": scores["helpfulness"]},
            {"indicator": "Complaint awareness", "score": scores["complaint_awareness"]},
            {"indicator": "Data completeness", "score": scores["completeness"]},
        ]
    ).dropna(subset=["score"])
    if not score_df.empty:
        fig = px.bar(
            score_df,
            x="indicator",
            y="score",
            template=template,
            title="Key Outcome Indicators",
            color="score",
            color_continuous_scale=scale,
            text=score_df["score"].map(lambda x: f"{x * 100:.1f}%"),
        )
        fig.update_yaxes(tickformat=".0%", range=[0, 1])
        fig.update_layout(showlegend=False, xaxis_title="Indicator", yaxis_title="Score")
        style_plotly_figure(fig, theme_mode)
        sections.append({"title": "Key outcome indicators", "caption": "Positive rates for the most relevant indicators available in the organization dataset.", "image": _figure_to_png_bytes(fig), "table": score_df.assign(score=score_df["score"].map(lambda x: f"{x * 100:.1f}%"))})

    distribution_col = _pick_distribution_column(data, indicators)
    if distribution_col:
        dist_df = _count_table(_safe_series(data, distribution_col), top_n=12, title_col=distribution_col)
        fig = px.bar(
            dist_df,
            x=distribution_col,
            y="count",
            template=template,
            title=f"Distribution of {distribution_col}",
            color="count",
            color_continuous_scale=scale,
        )
        fig.update_layout(showlegend=False, xaxis_title=distribution_col, yaxis_title="Interviews")
        style_plotly_figure(fig, theme_mode)
        sections.append({"title": f"{distribution_col} distribution", "caption": "Primary categorical breakdown for the selected organization.", "image": _figure_to_png_bytes(fig), "table": dist_df.copy()})

    return sections


def _build_chart_sections(
    scope_kind: str,
    scope_label: str,
    data: pd.DataFrame,
    progress: pd.DataFrame,
    indicators: dict[str, str | None],
    theme_mode: ThemeMode,
    scores: dict[str, float],
) -> list[dict[str, Any]]:
    if scope_kind == "organization":
        return _organization_chart_sections(scope_label, data, progress, indicators, theme_mode, scores)
    return _public_chart_sections(data, progress, indicators, theme_mode, scores)


def _ensure_docx_bidi(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "0")
    p_pr.append(bidi)


def _word_table(document: Document, title: str, df: pd.DataFrame) -> None:
    document.add_heading(_docx_safe_text(title), level=2)
    display = _display_df(df)
    if display.empty:
        document.add_paragraph("No data available.")
        return

    table = document.add_table(rows=1, cols=len(display.columns))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for idx, col in enumerate(display.columns):
        table.rows[0].cells[idx].text = _docx_safe_text(col)

    for _, row in display.iterrows():
        cells = table.add_row().cells
        for idx, value in enumerate(row.tolist()):
            cells[idx].text = _docx_safe_text(value)


def _register_pdf_font() -> str:
    for path in [Path("C:/Windows/Fonts/arial.ttf"), Path("C:/Windows/Fonts/segoeui.ttf"), Path("C:/Windows/Fonts/tahoma.ttf")]:
        if path.exists():
            try:
                pdfmetrics.registerFont(TTFont("ReportBaseFont", str(path)))
                return "ReportBaseFont"
            except Exception:
                continue
    return "Helvetica"


def _pdf_table(df: pd.DataFrame, font_name: str) -> Table:
    display = _display_df(df)
    if display.empty:
        display = pd.DataFrame([{"Info": "No data available."}])

    rows = [display.columns.tolist()] + display.values.tolist()
    table = Table(rows, repeatRows=1)
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1F4E79")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTNAME", (0, 0), (-1, -1), font_name),
                ("FONTSIZE", (0, 0), (-1, -1), 8),
                ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#C8D3DF")),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F5F7FA")]),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ]
        )
    )
    return table


def _validate_docx_bytes(payload: bytes) -> bool:
    if not payload:
        return False
    blob = io.BytesIO(payload)
    if not zipfile.is_zipfile(blob):
        return False
    try:
        Document(io.BytesIO(payload))
    except Exception:
        return False
    return True


def _build_word_fallback(bundle: dict[str, Any], reason: str) -> bytes:
    document = Document()
    document.add_heading(_docx_safe_text(bundle["title"]), level=0)
    document.add_paragraph(_docx_safe_text(bundle["subtitle"]))
    document.add_paragraph(_docx_safe_text(f"Generated at: {bundle['generated_at']}"))
    document.add_paragraph(_docx_safe_text(f"Scope: {bundle['scope_label']}"))
    document.add_paragraph(_docx_safe_text(f"Date coverage: {bundle['date_coverage']}"))
    document.add_paragraph(_docx_safe_text("Note: A safe fallback layout was used to preserve Word compatibility."))
    document.add_paragraph(_docx_safe_text(f"Builder detail: {reason}"))
    _word_table(document, "KPI Summary", bundle["summary_table"])
    _word_table(document, "Indicator Table", bundle["indicator_table"])
    _word_table(document, "Progress Table", bundle["progress_table"])
    _word_table(document, "Dataset Preview", bundle["dataset_preview"])
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def _build_word_bytes(bundle: dict[str, Any]) -> bytes:
    try:
        document = Document()
        document.core_properties.title = _docx_safe_text(bundle["title"])
        document.core_properties.subject = "Analytical monitoring report"
        document.core_properties.author = "UN Women Dashboard"

        normal = document.styles["Normal"]
        normal.font.name = "Calibri"
        normal.font.size = Pt(10.5)

        section = document.sections[0]
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

        title = document.add_heading(_docx_safe_text(bundle["title"]), level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle = document.add_paragraph(_docx_safe_text(bundle["subtitle"]))
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _ensure_docx_bidi(subtitle)

        document.add_paragraph(_docx_safe_text(f"Generated at: {bundle['generated_at']}"))
        document.add_paragraph(_docx_safe_text(f"Scope: {bundle['scope_label']}"))
        document.add_paragraph(_docx_safe_text(f"Date coverage: {bundle['date_coverage']}"))

        document.add_heading("Executive Summary", level=1)
        for item in bundle["insights"]:
            document.add_paragraph(_docx_safe_text(item), style="List Bullet")

        document.add_heading("Key Risks", level=1)
        for item in bundle["risks"]:
            document.add_paragraph(_docx_safe_text(item), style="List Bullet")

        document.add_heading("Recommendations", level=1)
        for item in bundle["recommendations"]:
            document.add_paragraph(_docx_safe_text(item), style="List Bullet")

        _word_table(document, "KPI Summary", bundle["summary_table"])
        _word_table(document, "Indicator Table", bundle["indicator_table"])

        export_sections = bundle["export_sections"][:MAX_WORD_CHART_SECTIONS]
        for idx, chart in enumerate(export_sections, start=1):
            if idx > 1:
                document.add_section(WD_SECTION.NEW_PAGE)
            document.add_heading(_docx_safe_text(chart["title"]), level=1)
            document.add_paragraph(_docx_safe_text(chart["caption"]))
            try:
                document.add_picture(io.BytesIO(chart["image"]), width=Inches(DOCX_IMAGE_WIDTH_INCH))
            except Exception:
                document.add_paragraph("Chart image could not be embedded in this section.")
            _word_table(document, f"{chart['title']} data", chart["table"])

        if len(bundle["export_sections"]) > MAX_WORD_CHART_SECTIONS:
            remaining = len(bundle["export_sections"]) - MAX_WORD_CHART_SECTIONS
            document.add_paragraph(_docx_safe_text(f"Note: {remaining} additional chart sections were skipped in Word export for file stability."))

        document.add_page_break()
        _word_table(document, "Progress Table", bundle["progress_table"])
        _word_table(document, "Dataset Preview", bundle["dataset_preview"])

        output = io.BytesIO()
        document.save(output)
        payload = output.getvalue()
        if _validate_docx_bytes(payload):
            return payload
        return _build_word_fallback(bundle, "Primary document validation failed.")
    except Exception as exc:
        return _build_word_fallback(bundle, f"Primary document build failed: {exc}")


def _build_pdf_bytes(bundle: dict[str, Any]) -> bytes:
    buffer = io.BytesIO()
    font_name = _register_pdf_font()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        leftMargin=0.55 * inch,
        rightMargin=0.55 * inch,
        topMargin=0.6 * inch,
        bottomMargin=0.6 * inch,
        title=bundle["title"],
        author="UN Women Dashboard",
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("ReportTitle", parent=styles["Title"], fontName=font_name, fontSize=20, leading=24, alignment=TA_LEFT, textColor=colors.HexColor("#123B63"), spaceAfter=10)
    body_style = ParagraphStyle("Body", parent=styles["BodyText"], fontName=font_name, fontSize=10, leading=14, spaceAfter=6)
    heading_style = ParagraphStyle("Heading", parent=styles["Heading2"], fontName=font_name, fontSize=14, leading=18, textColor=colors.HexColor("#123B63"), spaceBefore=10, spaceAfter=8)

    story: list[Any] = [
        Paragraph(bundle["title"], title_style),
        Paragraph(bundle["subtitle"], body_style),
        Paragraph(f"Generated at: {bundle['generated_at']}", body_style),
        Paragraph(f"Scope: {bundle['scope_label']}", body_style),
        Paragraph(f"Date coverage: {bundle['date_coverage']}", body_style),
        Spacer(1, 0.15 * inch),
        Paragraph("Executive Summary", heading_style),
    ]

    for item in bundle["insights"]:
        story.append(Paragraph(f"- {item}", body_style))
    story.append(Paragraph("Key Risks", heading_style))
    for item in bundle["risks"]:
        story.append(Paragraph(f"- {item}", body_style))
    story.append(Paragraph("Recommendations", heading_style))
    for item in bundle["recommendations"]:
        story.append(Paragraph(f"- {item}", body_style))

    story.extend([Paragraph("KPI Summary", heading_style), _pdf_table(bundle["summary_table"], font_name), Spacer(1, 0.12 * inch), Paragraph("Indicator Table", heading_style), _pdf_table(bundle["indicator_table"], font_name)])

    for chart in bundle["export_sections"]:
        story.extend([PageBreak(), Paragraph(chart["title"], heading_style), Paragraph(chart["caption"], body_style), Spacer(1, 0.08 * inch), Image(io.BytesIO(chart["image"]), width=6.9 * inch, height=4.1 * inch), Spacer(1, 0.12 * inch), _pdf_table(chart["table"], font_name)])

    story.extend([PageBreak(), Paragraph("Progress Table", heading_style), _pdf_table(bundle["progress_table"], font_name), Spacer(1, 0.12 * inch), Paragraph("Dataset Preview", heading_style), _pdf_table(bundle["dataset_preview"], font_name)])

    doc.build(story)
    return buffer.getvalue()


def _write_sheet(writer: pd.ExcelWriter, name: str, df: pd.DataFrame) -> None:
    frame = df.copy()
    for col in frame.columns:
        frame[col] = frame[col].map(_coerce_text)
    sheet_name = _safe_sheet_name(name)
    frame.to_excel(writer, sheet_name=sheet_name, index=False)
    worksheet = writer.sheets[sheet_name]
    for idx, col in enumerate(frame.columns):
        worksheet.set_column(idx, idx, min(max(len(str(col)) + 2, 16), 40))


def _build_excel_bytes(bundle: dict[str, Any]) -> bytes:
    buffer = io.BytesIO()
    with pd.ExcelWriter(buffer, engine="xlsxwriter") as writer:
        workbook = writer.book
        header_fmt = workbook.add_format({"bold": True, "bg_color": "#1F4E79", "font_color": "#FFFFFF", "border": 1})
        text_fmt = workbook.add_format({"text_wrap": True, "valign": "top", "border": 1})
        sheet_names: list[str] = []

        for name, df in [("summary", bundle["summary_table"]), ("indicators", bundle["indicator_table"]), ("progress", bundle["progress_table"]), ("dataset_preview", bundle["dataset_preview"]), ("dataset_full", bundle["dataset_export"])]:
            _write_sheet(writer, name, df)
            sheet_names.append(_safe_sheet_name(name))

        for idx, chart in enumerate(bundle["export_sections"], start=1):
            name = f"chart_{idx}"
            _write_sheet(writer, name, chart["table"])
            sheet_names.append(_safe_sheet_name(name))

        charts_sheet = workbook.add_worksheet("charts")
        writer.sheets["charts"] = charts_sheet
        charts_sheet.set_column(0, 2, 28)
        charts_sheet.write("A1", bundle["title"], workbook.add_format({"bold": True, "font_size": 16, "font_color": "#123B63"}))
        charts_sheet.write("A2", bundle["subtitle"])
        row_cursor = 4
        for chart in bundle["export_sections"]:
            charts_sheet.write(row_cursor, 0, chart["title"], workbook.add_format({"bold": True, "font_size": 13}))
            charts_sheet.write(row_cursor + 1, 0, chart["caption"])
            charts_sheet.insert_image(row_cursor + 2, 0, f"{_safe_slug(chart['title'])}.png", {"image_data": io.BytesIO(chart["image"]), "x_scale": 0.58, "y_scale": 0.58})
            row_cursor += 24

        for sheet_name in sheet_names:
            worksheet = writer.sheets[sheet_name]
            if worksheet.dim_rowmax is None or worksheet.dim_colmax is None:
                continue
            worksheet.autofilter(0, 0, worksheet.dim_rowmax, worksheet.dim_colmax)
            worksheet.freeze_panes(1, 0)
            worksheet.set_row(0, 22, header_fmt)
            worksheet.conditional_format(1, 0, worksheet.dim_rowmax, worksheet.dim_colmax, {"type": "no_blanks", "format": text_fmt})

    return buffer.getvalue()


@st.cache_data(show_spinner=False, ttl=600)
def build_report_artifacts(
    scope_kind: str,
    scope_label: str,
    data: pd.DataFrame,
    progress: pd.DataFrame,
    indicators: dict[str, str | None],
    theme_mode: ThemeMode,
) -> dict[str, Any]:
    scoped_progress = _relevant_progress(scope_kind, data, progress)
    date_column = _preferred_date_column(data, indicators)
    summary_items, indicator_df, scores = _build_kpi_summary(scope_kind, scope_label, data, scoped_progress, indicators)
    insights, risks, recommendations = _build_insights(scope_kind, scope_label, data, scoped_progress, indicators, scores)
    chart_sections = _build_chart_sections(scope_kind, scope_label, data, scoped_progress, indicators, theme_mode, scores)
    org_activity_sections = _org_activity_sections(data, date_column, theme_mode)
    _, _, date_label = _date_window(data, date_column)

    dataset_export = data.copy()
    bundle: dict[str, Any] = {
        "title": f"{'Organization' if scope_kind == 'organization' else 'Public'} Analytical Report",
        "subtitle": "Complete analytical package with visualizations, KPI interpretation, risk notes, and export-ready tables.",
        "generated_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "scope_kind": scope_kind,
        "scope_label": scope_label,
        "date_coverage": date_label,
        "summary_table": _summary_export_df(summary_items),
        "indicator_table": indicator_df.copy(),
        "insights": insights,
        "risks": risks,
        "recommendations": recommendations,
        "charts": chart_sections,
        "org_activity_sections": org_activity_sections,
        "export_sections": chart_sections + org_activity_sections,
        "progress_table": scoped_progress.copy(),
        "dataset_preview": _display_df(dataset_export, limit=MAX_REPORT_TABLE_ROWS),
        "dataset_export": dataset_export.copy(),
    }
    bundle["excel_bytes"] = _build_excel_bytes(bundle)
    bundle["word_bytes"] = _build_word_bytes(bundle)
    bundle["pdf_bytes"] = _build_pdf_bytes(bundle)
    bundle["file_stub"] = f"{_safe_slug(scope_label)}_{scope_kind}_analytical_report"
    return bundle

